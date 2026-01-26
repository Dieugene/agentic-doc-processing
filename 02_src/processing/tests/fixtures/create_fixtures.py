"""
Script to create test fixture files for Converter testing.

Run this script to generate sample.docx and sample.xlsx files.
Requires: python-docx, openpyxl
"""
from pathlib import Path


def create_sample_docx():
    """Create sample.docx fixture."""
    try:
        from docx import Document
    except ImportError:
        print("python-docx not installed. Install with: pip install python-docx")
        return

    doc = Document()

    # Title
    doc.add_heading("Тестовый документ DOCX", 0)

    # Regular paragraph
    doc.add_paragraph("Это пример текста для проверки конвертации DOCX в PDF.")

    # Bold paragraph
    p = doc.add_paragraph()
    run = p.add_run("Это жирный текст.")
    run.bold = True

    # Italic paragraph
    p = doc.add_paragraph()
    run = p.add_run("Это курсивный текст.")
    run.italic = True

    # Bold + Italic
    p = doc.add_paragraph()
    run = p.add_run("Это жирный курсив.")
    run.bold = True
    run.italic = True

    # Mixed content
    doc.add_paragraph("Разные размеры шрифта и стили форматирования.")

    # More content
    doc.add_paragraph("Второй параграф с обычным текстом.")
    doc.add_paragraph("Третий параграф для проверки переноса строк.")

    # Save
    output_path = Path(__file__).parent / "sample.docx"
    doc.save(output_path)
    print(f"Created {output_path}")


def create_sample_xlsx():
    """Create sample.xlsx fixture."""
    try:
        from openpyxl import Workbook
    except ImportError:
        print("openpyxl not installed. Install with: pip install openpyxl")
        return

    wb = Workbook()

    # First sheet
    ws1 = wb.active
    ws1.title = "Лист 1"

    # Headers
    headers = ["ID", "Название", "Количество", "Цена", "Сумма"]
    for col, header in enumerate(headers, 1):
        ws1.cell(row=1, column=col, value=header)

    # Data rows
    data = [
        [1, "Товар А", 10, 100.0, 1000.0],
        [2, "Товар Б", 5, 250.50, 1252.50],
        [3, "Товар В", 20, 50.0, 1000.0],
        [4, "Товар Г", 15, 75.0, 1125.0],
    ]
    for row_idx, row_data in enumerate(data, 2):
        for col_idx, value in enumerate(row_data, 1):
            ws1.cell(row=row_idx, column=col_idx, value=value)

    # Second sheet
    ws2 = wb.create_sheet("Лист 2")
    headers2 = ["Дата", "Сотрудник", "Часы", "Ставка", "Итого"]
    for col, header in enumerate(headers2, 1):
        ws2.cell(row=1, column=col, value=header)

    data2 = [
        ["2025-01-01", "Иванов", 8, 500.0, 4000.0],
        ["2025-01-02", "Петров", 6, 450.0, 2700.0],
        ["2025-01-03", "Сидоров", 8, 500.0, 4000.0],
    ]
    for row_idx, row_data in enumerate(data2, 2):
        for col_idx, value in enumerate(row_data, 1):
            ws2.cell(row=row_idx, column=col_idx, value=value)

    # Third sheet
    ws3 = wb.create_sheet("Пустой лист")
    # Empty sheet

    # Save
    output_path = Path(__file__).parent / "sample.xlsx"
    wb.save(output_path)
    print(f"Created {output_path}")


def create_sample_pdf():
    """Create sample.pdf fixture for Renderer testing."""
    try:
        from fpdf import FPDF
    except ImportError:
        print("fpdf2 not installed. Install with: pip install fpdf2")
        return

    pdf = FPDF()
    pdf.add_page()

    # Page 1 - Title and content
    pdf.set_font("Arial", size=12)
    pdf.cell(0, 10, txt="Test PDF Document - Page 1", ln=True, align="C")
    pdf.ln(5)

    pdf.set_font("Arial", size=10)
    pdf.multi_cell(0, 7, txt="This is a sample PDF document for testing the Renderer module.")
    pdf.ln(3)
    pdf.multi_cell(0, 7, txt="Page 1 contains basic text content for rendering to PNG.")
    pdf.ln(3)
    pdf.cell(0, 7, txt="Line 1: Some sample text")
    pdf.ln(7)
    pdf.cell(0, 7, txt="Line 2: More sample text")
    pdf.ln(7)
    pdf.cell(0, 7, txt="Line 3: Even more text")

    # Page 2 - More content
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(0, 10, txt="Test PDF Document - Page 2", ln=True, align="C")
    pdf.ln(5)

    pdf.set_font("Arial", size=10)
    pdf.multi_cell(0, 7, txt="This is the second page of the test PDF.")
    pdf.ln(3)
    pdf.cell(0, 7, txt="Different content on page 2")

    # Page 3 - Final page
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(0, 10, txt="Test PDF Document - Page 3", ln=True, align="C")
    pdf.ln(5)

    pdf.set_font("Arial", size=10)
    pdf.multi_cell(0, 7, txt="This is the third and final page of the test PDF.")
    pdf.ln(3)
    pdf.cell(0, 7, txt="End of test document")

    # Save
    output_path = Path(__file__).parent / "sample.pdf"
    pdf.output(str(output_path))
    print(f"Created {output_path}")


if __name__ == "__main__":
    print("Creating test fixtures...")
    create_sample_docx()
    create_sample_xlsx()
    create_sample_pdf()
    print("\nDone! Run this script from virtual environment with dependencies installed:")
    print("  pip install python-docx openpyxl fpdf2")
