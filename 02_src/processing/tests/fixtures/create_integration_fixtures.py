"""
Create integration test fixture files.

Run this script to generate sample.docx, sample.xlsx, sample.pdf
for integration testing of the full document processing pipeline.

Requires: python-docx, openpyxl, fpdf2
"""
from pathlib import Path


def create_sample_docx():
    """Create sample.docx with sections and a NUMERIC table."""
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        print("python-docx not installed. Install with: pip install python-docx")
        return

    doc = Document()

    # Title
    title = doc.add_heading("Test Document for Integration Testing", 0)

    # Section 1
    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph("This is the first section of the test document.")
    doc.add_paragraph("This describes the main purpose of the document.")

    # Section 1.1
    doc.add_heading("1.1. General Provisions", level=2)
    doc.add_paragraph("This is a subsection of the first section.")

    # Add NUMERIC table to section 1.1
    doc.add_paragraph("Table 1. Financial Indicators:")
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'

    # Header row
    headers = ["Indicator", "2023", "2024", "Change"]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Data rows
    data = [
        ["Revenue", "1000", "1200", "+20%"],
        ["Profit", "200", "250", "+25%"],
        ["Expenses", "800", "950", "+19%"],
        ["Margin", "20%", "21%", "+1%"],
    ]
    for row_idx, row_data in enumerate(data, start=1):
        for col_idx, value in enumerate(row_data):
            table.rows[row_idx].cells[col_idx].text = value

    # Section 1.2
    doc.add_heading("1.2. Details", level=2)
    doc.add_paragraph("Additional information in this subsection.")

    # Section 2
    doc.add_heading("2. Conclusion", level=1)
    doc.add_paragraph("Final section of the document.")
    doc.add_paragraph("Final conclusions and recommendations.")

    # Save
    output_path = Path(__file__).parent / "sample.docx"
    doc.save(output_path)
    print(f"Created {output_path}")


def create_sample_xlsx():
    """Create sample.xlsx with multiple sheets and NUMERIC tables."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, Alignment
    except ImportError:
        print("openpyxl not installed. Install with: pip install openpyxl")
        return

    wb = Workbook()

    # Sheet 1: Text description + NUMERIC table
    ws1 = wb.active
    ws1.title = "Sales Report"

    # Add some text (as comments in cells)
    ws1['A1'] = "Table 1. Sales by Region"
    ws1['A1'].font = Font(bold=True, size=12)

    # NUMERIC table with data
    headers = ["Region", "Q1", "Q2", "Q3", "Q4", "Total"]
    for col, header in enumerate(headers, 1):
        cell = ws1.cell(row=3, column=col, value=header)
        cell.font = Font(bold=True)

    sales_data = [
        ["Moscow", 1000, 1200, 1100, 1300, 4600],
        ["St. Petersburg", 800, 900, 850, 950, 3500],
        ["Kazan", 500, 600, 550, 700, 2350],
        ["Novosibirsk", 450, 500, 480, 600, 2030],
        ["Total", 2750, 3200, 2980, 3550, 12480],
    ]

    for row_idx, row_data in enumerate(sales_data, start=4):
        for col_idx, value in enumerate(row_data, 1):
            cell = ws1.cell(row=row_idx, column=col_idx, value=value)
            if row_idx == len(sales_data) + 3:  # Total row
                cell.font = Font(bold=True)
            cell.alignment = Alignment(horizontal='right')

    # Sheet 2: Only NUMERIC table
    ws2 = wb.create_sheet("Financial Indicators")

    ws2['A1'] = "Table 2. Financial Indicators (thousand RUB)"
    ws2['A1'].font = Font(bold=True, size=12)

    headers = ["Indicator", "2023", "2024", "Change %"]
    for col, header in enumerate(headers, 1):
        cell = ws2.cell(row=3, column=col, value=header)
        cell.font = Font(bold=True)

    financial_data = [
        ["Revenue", 50000, 60000, 20.0],
        ["Cost of Goods", 35000, 40000, 14.3],
        ["Gross Profit", 15000, 20000, 33.3],
        ["Operating Expenses", 8000, 9000, 12.5],
        ["Net Profit", 7000, 11000, 57.1],
    ]

    for row_idx, row_data in enumerate(financial_data, start=4):
        for col_idx, value in enumerate(row_data, 1):
            cell = ws2.cell(row=row_idx, column=col_idx, value=value)
            if isinstance(value, (int, float)):
                cell.alignment = Alignment(horizontal='right')

    # Save
    output_path = Path(__file__).parent / "sample.xlsx"
    wb.save(output_path)
    print(f"Created {output_path}")


def create_sample_pdf():
    """Create sample.pdf with TEXT_MATRIX tables."""
    try:
        from fpdf import FPDF
    except ImportError:
        print("fpdf2 not installed. Install with: pip install fpdf2")
        return

    pdf = FPDF()

    # Page 1
    pdf.add_page()
    pdf.set_font("Arial", 'B', 16)
    pdf.cell(0, 10, text="Document with Text Tables", ln=True, align="C")
    pdf.ln(5)

    pdf.set_font("Arial", size=12)
    pdf.cell(0, 10, text="1. Introduction", ln=True)
    pdf.ln(3)
    pdf.multi_cell(0, 7, text="This is a sample document with text tables for integration testing.")
    pdf.ln(5)

    # Page 2 - Section 1.1 with TEXT_MATRIX table
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(0, 10, text="1.1. Comparative Analysis", ln=True)
    pdf.ln(5)

    # TEXT_MATRIX table 1 - comparison table
    pdf.set_font("Arial", 'B', 11)
    pdf.cell(60, 7, text="Parameter", border=1)
    pdf.cell(60, 7, text="Option A", border=1)
    pdf.cell(60, 7, text="Option B", border=1, ln=True)
    pdf.set_font("Arial", size=10)

    table1_data = [
        ("Cost", "Lower", "Higher"),
        ("Speed", "Faster", "Slower"),
        ("Quality", "Good", "Excellent"),
        ("Reliability", "95%", "99%"),
    ]
    for row in table1_data:
        pdf.cell(60, 7, text=row[0], border=1)
        pdf.cell(60, 7, text=row[1], border=1)
        pdf.cell(60, 7, text=row[2], border=1, ln=True)

    pdf.ln(10)

    # Page 3 - Section 1.2 with another TEXT_MATRIX table
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(0, 10, text="1.2. Category Statistics", ln=True)
    pdf.ln(5)

    # TEXT_MATRIX table 2
    pdf.set_font("Arial", 'B', 11)
    pdf.cell(50, 7, text="Category", border=1)
    pdf.cell(50, 7, text="Count", border=1)
    pdf.cell(50, 7, text="Percent", border=1)
    pdf.cell(50, 7, text="Status", border=1, ln=True)
    pdf.set_font("Arial", size=10)

    table2_data = [
        ("Category 1", "150", "45%", "Active"),
        ("Category 2", "100", "30%", "In Review"),
        ("Category 3", "80", "25%", "Completed"),
    ]
    for row in table2_data:
        pdf.cell(50, 7, text=row[0], border=1)
        pdf.cell(50, 7, text=row[1], border=1)
        pdf.cell(50, 7, text=row[2], border=1)
        pdf.cell(50, 7, text=row[3], border=1, ln=True)

    # Page 4 - Section 2
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(0, 10, text="2. Conclusion", ln=True)
    pdf.ln(5)
    pdf.multi_cell(0, 7, text="This section contains the final conclusions of the document.")

    # Save
    output_path = Path(__file__).parent / "sample.pdf"
    pdf.output(str(output_path))
    print(f"Created {output_path}")


def create_expected_results():
    """Create expected_results JSON files for validation."""
    import json

    expected_results_dir = Path(__file__).parent / "expected_results"
    expected_results_dir.mkdir(exist_ok=True)

    # Expected result for DOCX
    docx_expected = {
        "expected_nodes": 6,  # root + 3 sections + 2 subsections
        "expected_hierarchy": {
            "root": ["section_1"],
            "section_1": ["section_1.1", "section_1.2"],
            "section_1.1": [],
            "section_1.2": [],
        },
        "expected_tables": [
            {"id": "table_1", "type": "NUMERIC", "attached_to": "section_1.1"}
        ]
    }

    with open(expected_results_dir / "docx_skeleton.json", "w", encoding="utf-8") as f:
        json.dump(docx_expected, f, ensure_ascii=False, indent=2)
    print(f"Created {expected_results_dir / 'docx_skeleton.json'}")

    # Expected result for XLSX
    xlsx_expected = {
        "expected_nodes": 3,  # root + 2 sheets (each sheet becomes a section)
        "expected_hierarchy": {
            "root": ["section_1"],
        },
        "expected_tables": [
            {"id": "table_1", "type": "NUMERIC", "attached_to": "section_1"}
        ]
    }

    with open(expected_results_dir / "xlsx_skeleton.json", "w", encoding="utf-8") as f:
        json.dump(xlsx_expected, f, ensure_ascii=False, indent=2)
    print(f"Created {expected_results_dir / 'xlsx_skeleton.json'}")

    # Expected result for PDF
    pdf_expected = {
        "expected_nodes": 5,  # root + 2 sections + 2 subsections
        "expected_hierarchy": {
            "root": ["section_1"],
            "section_1": ["section_1.1", "section_1.2"],
        },
        "expected_tables": [
            {"id": "table_1", "type": "TEXT_MATRIX", "attached_to": "section_1.1"},
            {"id": "table_2", "type": "TEXT_MATRIX", "attached_to": "section_1.2"}
        ]
    }

    with open(expected_results_dir / "pdf_skeleton.json", "w", encoding="utf-8") as f:
        json.dump(pdf_expected, f, ensure_ascii=False, indent=2)
    print(f"Created {expected_results_dir / 'pdf_skeleton.json'}")


if __name__ == "__main__":
    print("Creating integration test fixtures...")
    print()
    create_sample_docx()
    create_sample_xlsx()
    create_sample_pdf()
    print()
    create_expected_results()
    print()
    print("Done! Fixtures created for integration testing.")
    print("Run tests with: pytest -m integration")
